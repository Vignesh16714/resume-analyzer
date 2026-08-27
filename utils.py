"""
utils.py
--------
File validation and text-extraction helpers for uploaded resumes.
Supports PDF (via PyMuPDF / fitz) and DOCX (via python-docx).

Every public function raises a `ResumeParsingError` with a clear,
user-facing message on failure so app.py can catch one exception type
and show a friendly st.error() instead of crashing.
"""

import os
import re
from typing import Tuple

import fitz  # PyMuPDF
import docx  # python-docx

from config import ALLOWED_EXTENSIONS, MAX_FILE_SIZE_MB


class ResumeParsingError(Exception):
    """Raised for any recoverable problem with an uploaded resume file."""
    pass


def validate_file(filename: str, file_bytes: bytes) -> None:
    """
    Validates extension and size BEFORE we attempt to parse anything.
    Fails fast with a clear message rather than letting a parser throw
    a cryptic low-level error later.
    """
    ext = os.path.splitext(filename)[1].lower()
    if ext not in ALLOWED_EXTENSIONS:
        raise ResumeParsingError(
            f"Unsupported file type '{ext}'. Please upload a PDF or DOCX file."
        )

    size_mb = len(file_bytes) / (1024 * 1024)
    if size_mb > MAX_FILE_SIZE_MB:
        raise ResumeParsingError(
            f"File is too large ({size_mb:.1f} MB). Max allowed is {MAX_FILE_SIZE_MB} MB."
        )

    if size_mb == 0:
        raise ResumeParsingError("The uploaded file is empty.")


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extracts plain text from a PDF's bytes using PyMuPDF."""
    try:
        text_chunks = []
        with fitz.open(stream=file_bytes, filetype="pdf") as doc:
            if doc.page_count == 0:
                raise ResumeParsingError("The PDF has no pages.")
            for page in doc:
                text_chunks.append(page.get_text())
        text = "\n".join(text_chunks).strip()
    except ResumeParsingError:
        raise
    except Exception as exc:
        # Wrap any low-level PyMuPDF error in our own exception type.
        raise ResumeParsingError(f"Could not read PDF file: {exc}") from exc

    if not text:
        raise ResumeParsingError(
            "No extractable text found in this PDF. It may be a scanned "
            "image — try a text-based PDF or a DOCX file instead."
        )
    return text


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extracts plain text from a DOCX's bytes using python-docx."""
    import io

    try:
        document = docx.Document(io.BytesIO(file_bytes))
        paragraphs = [p.text for p in document.paragraphs if p.text.strip()]
        # Also pull text out of any tables (skills tables are common in resumes)
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text)
        text = "\n".join(paragraphs).strip()
    except Exception as exc:
        raise ResumeParsingError(f"Could not read DOCX file: {exc}") from exc

    if not text:
        raise ResumeParsingError("No extractable text found in this DOCX file.")
    return text


def extract_text(filename: str, file_bytes: bytes) -> str:
    """
    Dispatches to the correct extractor based on file extension.
    Call validate_file() first (app.py does this) so this function can
    assume the extension is one of ALLOWED_EXTENSIONS.
    """
    ext = os.path.splitext(filename)[1].lower()
    if ext == ".pdf":
        return extract_text_from_pdf(file_bytes)
    elif ext == ".docx":
        return extract_text_from_docx(file_bytes)
    raise ResumeParsingError(f"Unsupported file type '{ext}'.")


def clean_text(text: str) -> str:
    """Normalizes whitespace so downstream NLP/regex logic is reliable."""
    text = text.replace("\r", "\n")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def word_count(text: str) -> int:
    return len(re.findall(r"\b\w+\b", text))


def find_contact_info(text: str) -> Tuple[bool, bool]:
    """Returns (has_email, has_phone) using simple, robust regexes."""
    has_email = bool(re.search(r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}", text))
    has_phone = bool(re.search(r"(\+?\d{1,3}[\s-]?)?(\(?\d{3}\)?[\s.-]?)\d{3}[\s.-]?\d{4}", text))
    return has_email, has_phone
